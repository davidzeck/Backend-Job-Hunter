"""
CV document-export tests (roadmap Phase A #4): docgen golden tests (real
rendering, no mocks), the draft state machine, the curate task with mocked AI,
and the render task with mocked S3.
"""
import io
import uuid
from datetime import datetime, timezone

import pytest
from docx import Document
from sqlalchemy import select, text

from app.core.database import async_session_maker, engine
from app.core.docgen import render_docx, render_pdf
from app.models.company import Company
from app.models.cv_draft import (
    CVDraft,
    DRAFT_STATUS_APPROVED,
    DRAFT_STATUS_FAILED,
    DRAFT_STATUS_GENERATING,
    DRAFT_STATUS_RENDERED,
    DRAFT_STATUS_REVIEW,
    DRAFT_STATUS_SUPERSEDED,
)
from app.models.job import Job
from app.models.job_source import JobSource
from app.models.user_cv import UserCV, UPLOAD_STATUS_READY
from app.schemas.cv import CVStructure
from app.services.cv_draft_service import CVDraftService
from fastapi import HTTPException

SAMPLE_STRUCTURE = {
    "contact": {
        "name": "Jane Dev",
        "email": "jane@dev.ke",
        "phone": "+254700000000",
        "location": "Nairobi",
        "links": ["github.com/janedev"],
    },
    "summary": "Backend engineer with 5 years of Python and FastAPI.",
    "skills": [{"category": "Languages", "items": ["Python", "Go", "SQL"]}],
    "experience": [
        {
            "title": "Senior Engineer",
            "company": "Acme",
            "location": "Nairobi",
            "start": "Jan 2022",
            "end": "Present",
            "bullets": ["Built the alerts pipeline", "Cut p95 latency 40%"],
        }
    ],
    "education": [{"degree": "BSc CS", "institution": "UoN", "year": "2019"}],
    "certifications": ["AWS Solutions Architect Associate"],
}


# ── Factories ───────────────────────────────────────────────────────────

async def _user_id(email):
    async with engine.begin() as db:
        r = await db.execute(text("SELECT id FROM users WHERE email=:e"), {"e": email})
        return r.scalar()


async def _make_cv_and_job(user_id, *, cv_status=UPLOAD_STATUS_READY,
                           parsed_structure=None, description="A real job description."):
    async with async_session_maker() as db:
        company = Company(name="Acme", slug=f"acme-{uuid.uuid4().hex[:6]}")
        db.add(company)
        await db.flush()
        source = JobSource(
            company_id=company.id, source_type="api",
            url="https://acme.example.com", scraper_class="greenhouse",
        )
        db.add(source)
        await db.flush()
        job = Job(
            source_id=source.id, company_id=company.id,
            external_id=f"e-{uuid.uuid4().hex[:8]}", title="Backend Engineer",
            description=description, apply_url="https://acme.example.com/1",
            discovered_at=datetime.now(timezone.utc), validation_status="valid",
        )
        cv = UserCV(
            user_id=user_id, filename="cv.pdf", file_path="",
            upload_status=cv_status, full_text="Jane Dev. Python engineer.",
            parsed_structure=parsed_structure,
        )
        db.add_all([job, cv])
        await db.commit()
        return cv.id, job.id


async def _draft_row(draft_id):
    async with async_session_maker() as db:
        return (await db.execute(
            select(CVDraft).where(CVDraft.id == draft_id)
        )).scalar_one()


# ── docgen golden tests (real rendering, no mocks) ──────────────────────

def test_render_docx_structure_and_no_tables():
    data = render_docx(SAMPLE_STRUCTURE)
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]

    assert "Jane Dev" in texts[0]
    for heading in ("SUMMARY", "SKILLS", "WORK EXPERIENCE", "EDUCATION", "CERTIFICATIONS"):
        assert heading in texts
    assert any(t.startswith("- Built the alerts pipeline") for t in texts)
    assert "Languages: Python, Go, SQL" in texts
    assert len(doc.tables) == 0  # ATS rule: no tables


def test_render_pdf_text_extractable():
    import pdfplumber

    data = render_pdf(SAMPLE_STRUCTURE)
    assert data[:5] == b"%PDF-"
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        assert len(pdf.pages) >= 1
        full = "\n".join(page.extract_text() or "" for page in pdf.pages)
    assert "Jane Dev" in full
    assert "WORK EXPERIENCE" in full
    assert "Built the alerts pipeline" in full


def test_render_handles_empty_sections():
    minimal = {"contact": {"name": "X"}, "summary": "", "skills": [],
               "experience": [], "education": [], "certifications": []}
    assert len(render_docx(minimal)) > 0
    assert render_pdf(minimal)[:5] == b"%PDF-"


# ── Draft state machine (service level) ─────────────────────────────────

async def test_curate_supersedes_live_draft(registered_user):
    uid = await _user_id(registered_user["email"])
    cv_id, job_id = await _make_cv_and_job(uid)
    svc = CVDraftService()

    async with async_session_maker() as db:
        first = await svc.start_curate(db, uid, cv_id, job_id)
    async with async_session_maker() as db:
        second = await svc.start_curate(db, uid, cv_id, job_id)

    first_row = await _draft_row(uuid.UUID(first["draft_id"]))
    second_row = await _draft_row(uuid.UUID(second["draft_id"]))
    assert first_row.status == DRAFT_STATUS_SUPERSEDED
    assert second_row.status == DRAFT_STATUS_GENERATING


async def test_curate_requires_ready_cv(registered_user):
    uid = await _user_id(registered_user["email"])
    cv_id, job_id = await _make_cv_and_job(uid, cv_status="processing")
    svc = CVDraftService()
    async with async_session_maker() as db:
        with pytest.raises(HTTPException) as exc:
            await svc.start_curate(db, uid, cv_id, job_id)
    assert exc.value.status_code == 409


async def test_patch_only_in_review_and_approve_flow(registered_user):
    uid = await _user_id(registered_user["email"])
    cv_id, job_id = await _make_cv_and_job(uid)
    svc = CVDraftService()

    async with async_session_maker() as db:
        started = await svc.start_curate(db, uid, cv_id, job_id)
    draft_id = uuid.UUID(started["draft_id"])

    edited = CVStructure.model_validate(SAMPLE_STRUCTURE)

    # PATCH while still generating → 409
    async with async_session_maker() as db:
        with pytest.raises(HTTPException) as exc:
            await svc.update_draft(db, uid, draft_id, edited)
    assert exc.value.status_code == 409

    # Approve while still generating → 409
    async with async_session_maker() as db:
        with pytest.raises(HTTPException) as exc:
            await svc.approve_draft(db, uid, draft_id)
    assert exc.value.status_code == 409

    # Simulate the curate task finishing → review
    async with async_session_maker() as db:
        row = (await db.execute(select(CVDraft).where(CVDraft.id == draft_id))).scalar_one()
        row.content = {"original": SAMPLE_STRUCTURE, "tailored": SAMPLE_STRUCTURE,
                       "keywords_injected": ["FastAPI"]}
        row.status = DRAFT_STATUS_REVIEW
        await db.commit()

    # PATCH now works and persists
    async with async_session_maker() as db:
        resp = await svc.update_draft(db, uid, draft_id, edited)
    assert resp.content["tailored"]["contact"]["name"] == "Jane Dev"

    # Approve → approved (+ enqueues render)
    async with async_session_maker() as db:
        await svc.approve_draft(db, uid, draft_id)
    assert (await _draft_row(draft_id)).status == DRAFT_STATUS_APPROVED

    # Second approve → 409 (double-approve guard)
    async with async_session_maker() as db:
        with pytest.raises(HTTPException) as exc:
            await svc.approve_draft(db, uid, draft_id)
    assert exc.value.status_code == 409

    # Download before rendered → 409
    async with async_session_maker() as db:
        with pytest.raises(HTTPException) as exc:
            await svc.get_download_url(db, uid, draft_id, "pdf")
    assert exc.value.status_code == 409


async def test_draft_ownership_404(registered_user):
    uid = await _user_id(registered_user["email"])
    cv_id, job_id = await _make_cv_and_job(uid)
    svc = CVDraftService()
    async with async_session_maker() as db:
        started = await svc.start_curate(db, uid, cv_id, job_id)
    stranger = uuid.uuid4()
    async with async_session_maker() as db:
        with pytest.raises(HTTPException) as exc:
            await svc.get_draft(db, stranger, uuid.UUID(started["draft_id"]))
    assert exc.value.status_code == 404


# ── Curate task (mocked AI) ─────────────────────────────────────────────

async def _make_generating_draft(uid, cv_id, job_id):
    async with async_session_maker() as db:
        draft = CVDraft(cv_id=cv_id, job_id=job_id, user_id=uid,
                        status=DRAFT_STATUS_GENERATING)
        db.add(draft)
        await db.commit()
        return draft.id


async def test_curate_task_happy_path_caches_parse(registered_user, monkeypatch):
    from app.core import ai
    from app.workers.tasks import _curate_cv

    uid = await _user_id(registered_user["email"])
    cv_id, job_id = await _make_cv_and_job(uid)  # parsed_structure=None
    draft_id = await _make_generating_draft(uid, cv_id, job_id)

    calls = {"parse": 0}

    async def fake_parse(cv_text):
        calls["parse"] += 1
        return SAMPLE_STRUCTURE

    async def fake_keywords(jd):
        return {"required_skills": ["FastAPI"]}

    async def fake_analyze(cv_text, jd, kw):
        return {"missing_keywords": ["FastAPI"]}

    async def fake_tailor(structure, jd, missing):
        assert missing == ["FastAPI"]
        return {"tailored": SAMPLE_STRUCTURE, "keywords_injected": ["FastAPI"]}

    monkeypatch.setattr(ai, "parse_cv_structure", fake_parse)
    monkeypatch.setattr(ai, "extract_keywords_from_jd", fake_keywords)
    monkeypatch.setattr(ai, "analyze_cv_against_jd", fake_analyze)
    monkeypatch.setattr(ai, "tailor_cv_full", fake_tailor)

    result = await _curate_cv(str(uid), str(cv_id), str(job_id), str(draft_id))
    assert result["status"] == DRAFT_STATUS_REVIEW
    assert calls["parse"] == 1

    row = await _draft_row(draft_id)
    assert row.status == DRAFT_STATUS_REVIEW
    assert row.content["keywords_injected"] == ["FastAPI"]
    assert row.content["original"]["contact"]["name"] == "Jane Dev"

    # Parse result cached on the CV row
    async with async_session_maker() as db:
        cv = (await db.execute(select(UserCV).where(UserCV.id == cv_id))).scalar_one()
        assert cv.parsed_structure["contact"]["name"] == "Jane Dev"

    # Second curate for the same CV reuses the cached parse (no new parse call)
    draft2 = await _make_generating_draft(uid, cv_id, job_id)
    result2 = await _curate_cv(str(uid), str(cv_id), str(job_id), str(draft2))
    assert result2["status"] == DRAFT_STATUS_REVIEW
    assert calls["parse"] == 1


async def test_curate_task_empty_ai_output_fails_draft(registered_user, monkeypatch):
    from app.core import ai
    from app.workers.tasks import _curate_cv

    uid = await _user_id(registered_user["email"])
    cv_id, job_id = await _make_cv_and_job(uid, parsed_structure=SAMPLE_STRUCTURE)
    draft_id = await _make_generating_draft(uid, cv_id, job_id)

    async def fake_keywords(jd):
        return {}

    async def fake_analyze(cv_text, jd, kw):
        return {"missing_keywords": []}

    async def fake_tailor(structure, jd, missing):
        return {"tailored": {}, "keywords_injected": []}  # effectively empty

    monkeypatch.setattr(ai, "extract_keywords_from_jd", fake_keywords)
    monkeypatch.setattr(ai, "analyze_cv_against_jd", fake_analyze)
    monkeypatch.setattr(ai, "tailor_cv_full", fake_tailor)

    result = await _curate_cv(str(uid), str(cv_id), str(job_id), str(draft_id))
    assert "error" in result
    row = await _draft_row(draft_id)
    assert row.status == DRAFT_STATUS_FAILED
    assert row.error


# ── Render task (mocked S3) ─────────────────────────────────────────────

async def test_generate_document_renders_and_marks(registered_user, monkeypatch):
    from app.core import storage
    from app.workers.tasks import _generate_cv_document

    uid = await _user_id(registered_user["email"])
    cv_id, job_id = await _make_cv_and_job(uid)
    async with async_session_maker() as db:
        draft = CVDraft(
            cv_id=cv_id, job_id=job_id, user_id=uid,
            status=DRAFT_STATUS_APPROVED,
            approved_at=datetime.now(timezone.utc),
            content={"original": SAMPLE_STRUCTURE, "tailored": SAMPLE_STRUCTURE,
                     "keywords_injected": []},
        )
        db.add(draft)
        await db.commit()
        draft_id = draft.id

    uploaded = {}

    async def fake_upload(s3_key, data, content_type):
        uploaded[s3_key] = (len(data), content_type)

    monkeypatch.setattr(storage, "upload_bytes", fake_upload)

    result = await _generate_cv_document(str(draft_id))
    assert result["status"] == DRAFT_STATUS_RENDERED

    row = await _draft_row(draft_id)
    assert row.status == DRAFT_STATUS_RENDERED
    assert row.docx_s3_key and row.docx_s3_key.endswith("cv.docx")
    assert row.pdf_s3_key and row.pdf_s3_key.endswith("cv.pdf")
    assert f"generated/{draft_id}" in row.docx_s3_key
    assert row.docx_s3_key in uploaded and row.pdf_s3_key in uploaded


async def test_generate_document_requires_approved(registered_user):
    from app.workers.tasks import _generate_cv_document

    uid = await _user_id(registered_user["email"])
    cv_id, job_id = await _make_cv_and_job(uid)
    draft_id = await _make_generating_draft(uid, cv_id, job_id)

    result = await _generate_cv_document(str(draft_id))
    assert "error" in result
    assert (await _draft_row(draft_id)).status == DRAFT_STATUS_GENERATING  # untouched