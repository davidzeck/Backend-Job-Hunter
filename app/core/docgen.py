"""
CV document rendering — CVStructure JSON → DOCX / PDF bytes.

One deliberately boring, ATS-safe template for both formats: single column,
standard section headings, Helvetica/Calibri, plain "- " bullets, no tables,
text boxes, images, headers, or footers — exactly the shapes ATS parsers
reward. Pure functions (bytes in-memory, no I/O) so tests can assert on real
rendered output without mocks.

Both renderers take the plain dict form of app/schemas/cv.py CVStructure;
empty sections are omitted.
"""
import io
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt
from fpdf import FPDF

_SECTION_ORDER = (
    ("summary", "SUMMARY"),
    ("skills", "SKILLS"),
    ("experience", "WORK EXPERIENCE"),
    ("education", "EDUCATION"),
    ("certifications", "CERTIFICATIONS"),
)


def _contact_line(contact: Dict[str, Any]) -> str:
    parts = [
        contact.get("email", ""),
        contact.get("phone", ""),
        contact.get("location", ""),
        *contact.get("links", []),
    ]
    return " | ".join(p for p in parts if p)


def _experience_heading(exp: Dict[str, Any]) -> str:
    left = " — ".join(p for p in (exp.get("title", ""), exp.get("company", "")) if p)
    dates = " – ".join(p for p in (exp.get("start", ""), exp.get("end", "")) if p)
    loc = exp.get("location", "")
    right = ", ".join(p for p in (loc, dates) if p)
    return f"{left} ({right})" if right else left


def _education_line(edu: Dict[str, Any]) -> str:
    return " — ".join(
        p for p in (edu.get("degree", ""), edu.get("institution", ""), edu.get("year", "")) if p
    )


def _skill_lines(groups: List[Dict[str, Any]]) -> List[str]:
    lines = []
    for group in groups:
        items = ", ".join(group.get("items", []))
        if not items:
            continue
        category = group.get("category", "")
        lines.append(f"{category}: {items}" if category else items)
    return lines


# ── DOCX ──────────────────────────────────────────────────────────────────────

def render_docx(structure: Dict[str, Any]) -> bytes:
    """Render a CVStructure dict to DOCX bytes (python-docx)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    contact = structure.get("contact", {}) or {}
    name = contact.get("name", "")
    if name:
        para = doc.add_paragraph()
        run = para.add_run(name)
        run.bold = True
        run.font.size = Pt(16)
    line = _contact_line(contact)
    if line:
        doc.add_paragraph(line)

    def heading(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)

    for key, title in _SECTION_ORDER:
        value = structure.get(key)
        if not value:
            continue
        heading(title)
        if key == "summary":
            doc.add_paragraph(value)
        elif key == "skills":
            for skill_line in _skill_lines(value):
                doc.add_paragraph(skill_line)
        elif key == "experience":
            for exp in value:
                head = _experience_heading(exp)
                if head:
                    para = doc.add_paragraph()
                    para.add_run(head).bold = True
                for bullet in exp.get("bullets", []):
                    doc.add_paragraph(f"- {bullet}")
        elif key == "education":
            for edu in value:
                edu_line = _education_line(edu)
                if edu_line:
                    doc.add_paragraph(edu_line)
        elif key == "certifications":
            for cert in value:
                doc.add_paragraph(f"- {cert}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────────────

def _pdf_safe(text: str) -> str:
    """fpdf2 core fonts are latin-1 — replace anything outside it."""
    return text.encode("latin-1", errors="replace").decode("latin-1")


def render_pdf(structure: Dict[str, Any]) -> bytes:
    """Render a CVStructure dict to PDF bytes (fpdf2, core Helvetica)."""
    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(left=18, top=18, right=18)
    pdf.add_page()
    width = pdf.w - pdf.l_margin - pdf.r_margin

    def block(text: str, size: float = 10.5, bold: bool = False, pad: float = 1.0):
        pdf.set_font("Helvetica", "B" if bold else "", size)
        pdf.multi_cell(width, size * 0.55, _pdf_safe(text))
        pdf.ln(pad)

    contact = structure.get("contact", {}) or {}
    if contact.get("name"):
        block(contact["name"], size=16, bold=True)
    line = _contact_line(contact)
    if line:
        block(line)

    for key, title in _SECTION_ORDER:
        value = structure.get(key)
        if not value:
            continue
        pdf.ln(2)
        block(title, size=12, bold=True)
        if key == "summary":
            block(value)
        elif key == "skills":
            for skill_line in _skill_lines(value):
                block(skill_line)
        elif key == "experience":
            for exp in value:
                head = _experience_heading(exp)
                if head:
                    block(head, bold=True)
                for bullet in exp.get("bullets", []):
                    block(f"- {bullet}")
        elif key == "education":
            for edu in value:
                edu_line = _education_line(edu)
                if edu_line:
                    block(edu_line)
        elif key == "certifications":
            for cert in value:
                block(f"- {cert}")

    return bytes(pdf.output())
